import tkinter as tk
from tkinter.filedialog import askopenfilename
from docx import Document
from docx.enum.style import WD_STYLE_TYPE


def get_file_namepath():
	""" Open modal dialog to get file location from user
	and return file name and file path"""
	root = tk.Tk()
	root.withdraw()
	filepath = askopenfilename()
	name = filepath.split('/').pop()
	path = filepath.replace(name, '')
	return name, path, filepath


def process_doc(filename):
	""" Crawl through the file and add html tags """
	document = Document(filename)
	paragraphs = document.paragraphs
	tables = document.tables
	
	process_paragraph(paragraphs)
	for t in tables:
		rows = t.rows
		cols = t.columns
		for row in range(len(rows)):
			for col in range(len(cols)):
				cl = t.cell(row, col)
				cell_par = cl.paragraphs
				process_paragraph(cell_par)

	return document


def process_paragraph(parlist):
	for par in parlist:
		runs = par.runs
		for r in runs:
			update_fonts(r)
			update_style(r)


def update_fonts(r):
	""" Remove strikethrough text
		add subscript and superscript html tags """
	fonts = r.font
	if fonts.strike:
		text = r.text.replace(r.text, '')
		r.text = text
	if fonts.subscript:
		text = r.text.replace(r.text, '<sub>'+r.text+'</sub>')
		r.text = text
	if fonts.superscript:
		text = r.text.replace(r.text, '<sup>'+r.text+'</sup>')
		r.text = text


def update_style(r):
	""" Add HTML tags for text style """
	if r.underline:
		text = r.text.replace(r.text, '<u>'+r.text+'</u>')
		r.text = text
	if r.bold and r.text.count('<strong>') == 0 and r.text.count('</strong>') == 0:
		text = r.text.replace(r.text, '<strong>'+r.text+'</strong>')
		r.text = text
	if r.italic:
		text = r.text.replace(r.text, '<em>'+r.text+'</em>')
		r.text = text


def main():
	fname, fpath, fp = get_file_namepath()
	doc = process_doc(fp)

	# save updated document to same directory as original file
	doc.save(fpath + 'PROCESSED_' + fname)


if __name__ == '__main__':
	main()
